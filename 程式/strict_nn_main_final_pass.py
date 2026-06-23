from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def find_report() -> Path:
    candidates = [
        p
        for p in Path(".").glob("*.docx")
        if p.name.startswith("11127004_")
        and "Stroke_Prediction" in p.name
        and "NN主軸版" in p.name
        and p.stat().st_size > 100_000
    ]
    return max(candidates, key=lambda p: p.stat().st_mtime)


def style_run(run, size=11, bold=False):
    run.font.name = "Microsoft JhengHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph(paragraph, text, size=11, bold=False):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold)
    return paragraph


def insert_paragraph_after(paragraph, text, style=None, size=11, bold=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p = new_p
    inserted._element = new_p
    if style:
        inserted.style = style
    run = inserted.add_run(text)
    style_run(run, size=size, bold=bold)
    return inserted


def main():
    path = find_report()
    doc = Document(path)

    # Strengthen opening and purpose so the midterm-extension and NN-main roles
    # are both unmistakable.
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("本研究延伸先前以 Stroke Prediction Dataset"):
            set_paragraph(
                p,
                "本研究延續期中 Stroke Prediction 分析。期中階段已完成資料清理、Missing Value 處理、"
                "視覺化、Encoding / Normalization，以及 Logistic Regression、Random Forest、KNN 等傳統機器學習模型比較；"
                "期末則在此基礎上，依課程要求加入神經網路，並以 MLPClassifier 作為主要模型，"
                "檢驗非線性模型在不平衡中風資料上的風險排序能力。Logistic Regression 僅作為 benchmark / baseline / 可解釋基準模型，"
                "用來比較 MLP 的優勢與限制。",
            )
        elif t.startswith("本研究的第一個目的，是在期中分析基礎上建立 MLP Neural Network"):
            set_paragraph(
                p,
                "本研究的第一個目的，是在期中分析基礎上建立 MLP Neural Network，檢驗非線性模型是否能改善中風風險排序。"
                "第二個目的，是透過 GridSearchCV 系統化搜尋 MLP 架構、activation、regularization 與 learning rate。"
                "第三個目的，是在不使用測試集資訊的前提下，以 OOF F2-score 調整分類門檻，使模型更符合中風早期風險篩檢對 Recall 的需求。"
                "第四個目的，是保留 LR benchmark，確認 MLP 的改善屬於 PR-AUC / 風險排序能力，而非宣稱神經網路在所有指標上全面勝出。",
            )
        elif t.startswith("數值特徵包含 age、hypertension"):
            set_paragraph(
                p,
                "數值特徵包含 age、hypertension、heart_disease、avg_glucose_level 與 bmi。數值缺失值以中位數填補，再使用 StandardScaler 標準化。"
                "類別特徵包含 gender、ever_married、work_type、Residence_type 與 smoking_status；類別缺失值以眾數填補，再使用 One-Hot Encoding。"
                "所有前處理皆整合於 Pipeline 中，並分別套用於 MLP 主模型與 Logistic Regression benchmark，"
                "確保每一折交叉驗證都只使用該折訓練資料估計補值與標準化參數，降低資料洩漏風險。",
            )
        elif t == "9.3 與期中推薦模型的關係":
            set_paragraph(p, "9.2 與期中推薦模型的關係", size=13, bold=True)
        elif t == "9.4 Neural Network 主模型比較":
            set_paragraph(p, "9.3 Neural Network 主模型比較", size=13, bold=True)
        elif t.startswith("期中報告推薦 Balanced Logistic Regression"):
            set_paragraph(
                p,
                "期中報告推薦 Balanced Logistic Regression，其測試結果為 TP=40、FN=10、FP=251，Recall 為 0.80。"
                "期末 MLP 的測試結果為 TP=39、FN=11、FP=215，Recall 為 0.78。"
                "與期中相比，MLP 減少 36 個 False Positive，但多 1 個 False Negative；同時，MLP 的 PR-AUC 0.2736 高於 LR benchmark 的 0.2517。"
                "需注意，本研究中 LR 有兩個用途：期中 Balanced LR 用於回顧原始基線；期末 LR benchmark 則使用與 MLP 相同流程進行 threshold tuning，"
                "因此 FP 數值會不同，期中 Balanced LR 為 FP=251，期末 LR benchmark 為 FP=209。"
                "因此，期末神經網路的價值主要在於風險排序能力提升，而不是在所有分類指標上完全勝出。",
            )
        elif t.startswith("MLP 使用與 Logistic Regression 相同的資料切分與前處理 Pipeline"):
            set_paragraph(
                p,
                "MLP 使用與 LR benchmark 相同的資料切分與前處理 Pipeline，包含中位數補值、One-Hot Encoding 與 StandardScaler，"
                "以確保兩者比較基礎一致。神經網路搜尋的最佳設定為 hidden_layer_sizes=(64,)、activation='relu'、alpha=0.0001、"
                "learning_rate_init=0.001，並使用 early stopping 降低過度擬合風險。由於中風類別樣本較少，訓練時同樣使用正類權重較高的 sample_weight，"
                "使模型更重視中風個案。",
            )
        elif t.startswith("本研究完成期中報告的延伸"):
            set_paragraph(
                p,
                "本研究延續期中 Stroke Prediction 分析，期中以傳統機器學習模型建立基線，期末則加入 MLP 神經網路作為主模型。"
                "MLP GridSearchCV 找到的最佳參數為 hidden_layer_sizes=(64,)、activation='relu'、alpha=0.0001、learning_rate_init=0.001；"
                "接著利用訓練集 OOF 預測與 F2-score 選出約 0.245 的分類門檻。",
            )
        elif t.startswith("整體而言，若目標是早期風險篩檢"):
            set_paragraph(
                p,
                "整體而言，MLP 經 GridSearchCV 與 OOF F2 threshold tuning 後，在 PR-AUC 上高於 LR benchmark，顯示其風險排序能力有所提升；"
                "但 Recall 與 F2 略低於 LR，因此本研究不主張神經網路全面勝出，而是將它定位為期末神經網路延伸主軸，並誠實呈現其優勢與限制。"
                "未來研究應加入外部驗證、更多臨床特徵、成本敏感分析與模型校準，並由實際醫療情境決定可接受的 Recall 與 Precision 平衡。",
            )

    # Tighten table title/caption if present.
    for p in doc.paragraphs:
        if p.text.strip() == "表 2　Logistic Regression 與 MLP Neural Network 的測試集比較":
            set_paragraph(p, "表 2　MLP 主模型與 LR benchmark 的測試集比較", size=10, bold=True)

    doc.save(path)
    print(f"Updated {path}")


if __name__ == "__main__":
    main()
