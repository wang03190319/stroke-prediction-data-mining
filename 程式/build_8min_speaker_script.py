from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("11127004_王俊詠_Stroke_Prediction_8分鐘講稿.docx")

NAVY = "20354A"
RED = "E4473D"
TEAL = "2A8C91"
MUTED = "687786"
INK = "1C2936"
PALE = "EEF2F5"


SLIDES = [
    (
        "第 1 頁｜封面",
        "0:00-0:20",
        "各位老師、同學大家好，我是王俊詠，學號 11127004。今天報告的題目是「神經網路中風預測分析」。這份期末報告是延伸我的期中研究，但這次主軸不是只做 Logistic Regression，而是加入 MLP 神經網路，並用期中模型作為比較基準，看看非線性模型在中風預測上能不能帶來新的改善。",
    ),
    (
        "第 2 頁｜期中到期末的延伸",
        "0:20-0:50",
        "期中我已經完成 Dataset、Missing value、相關性分析、One-Hot Encoding、標準化與多個分類模型比較。期中結果顯示，Logistic Regression 搭配 class weight 可以達到不錯的 Recall。期末則把重點放在神經網路：使用 MLPClassifier 作為主模型，並透過 GridSearchCV 搜尋神經網路參數，再用 OOF F2 選擇較適合醫療預警的分類門檻。",
    ),
    (
        "第 3 頁｜Dataset",
        "0:50-1:20",
        "本研究使用 Healthcare Stroke Dataset，共 5,110 筆資料，其中中風個案只有 249 筆，占 4.87%。這代表資料高度不平衡。如果只看 Accuracy，模型可能全部預測未中風也會看起來很高分，但這對醫療預警沒有意義。因此本研究會更重視 Recall、F2-score、PR-AUC 和混淆矩陣。",
    ),
    (
        "第 4 頁｜視覺化",
        "1:20-1:50",
        "從視覺化可以看到，中風樣本比例很低，這也是模型訓練最困難的地方。年齡較高、平均血糖較高、具有高血壓或心臟病紀錄的族群，中風比例相對較高。不過這些特徵並不是單獨決定中風，而是多個風險因素共同累積，所以期末才嘗試用 MLP 學習特徵之間的非線性組合。",
    ),
    (
        "第 5 頁｜Missing value 與相關性",
        "1:50-2:20",
        "資料中主要缺失欄位是 BMI，共 201 筆。期末流程使用中位數補值，因為中位數較不容易被極端值影響。相關性熱圖顯示，年齡與中風的線性關係最明顯，血糖、高血壓與心臟病也有正相關。但相關性不等於因果，而且類別欄位需要 One-Hot Encoding 後才能進入模型。",
    ),
    (
        "第 6 頁｜編碼轉換與標準化",
        "2:20-2:50",
        "前處理使用 ColumnTransformer 和 Pipeline。數值欄位先補值再 StandardScaler，類別欄位先補眾數再 One-Hot Encoding。這樣做有兩個好處：第一，神經網路對特徵尺度敏感，標準化可以讓訓練更穩定；第二，所有前處理都包在 Pipeline 裡，交叉驗證每一折都只會學該折的訓練資料，避免資料洩漏。",
    ),
    (
        "第 7 頁｜神經網路流程",
        "2:50-3:20",
        "整體流程是先切分訓練集與測試集，測試集只在最後一次使用。訓練階段先完成 Pipeline，再用 GridSearchCV 搜尋 MLP 的 hidden layer、activation、alpha 和 learning rate。接著使用訓練集的 Out-of-Fold 預測來找 F2 最大的門檻，最後才拿固定好的模型和門檻到測試集評估。",
    ),
    (
        "第 8 頁｜MLP 模型設計",
        "3:20-3:55",
        "這份資料是表格型資料，不是影像，也不是時間序列，所以我選擇 MLP，也就是多層感知器，而不是 CNN 或 RNN。MLP 的好處是可以學習欄位之間的非線性關係，例如年齡、血糖、BMI、工作型態與婚姻狀態可能共同形成風險。期末主軸就是用 MLP 作為主要模型，LR 則保留為可解釋的比較基準。",
    ),
    (
        "第 9 頁｜MLP GridSearchCV",
        "3:55-4:25",
        "在神經網路參數搜尋中，我比較 hidden layer、activation、alpha 和 learning rate。最佳設定是 hidden_layer_sizes 等於 (64,)、activation 使用 ReLU、alpha 是 0.0001、learning_rate_init 是 0.001。交叉驗證的 scoring 使用 PR-AUC，因為資料不平衡時，PR-AUC 比 Accuracy 更能反映少數類別的排序能力。",
    ),
    (
        "第 10 頁｜Decision Threshold",
        "4:25-4:55",
        "神經網路輸出的是機率，如果直接用 0.5 當門檻，會漏掉很多中風個案。以預設門檻來看，False Negative 是 33，只抓到 17 位中風者。改用 OOF F2 選出的 0.245 門檻後，False Negative 降到 11，True Positive 提升到 39。這符合醫療預警的目標：寧可多一些誤報，也要減少漏判。",
    ),
    (
        "第 11 頁｜OOF F2 門檻",
        "4:55-5:25",
        "這裡特別強調，門檻不是用測試集調出來的，而是用訓練資料的 Out-of-Fold 預測決定。這樣可以避免把測試集拿來調參造成分數高估。最後選出的 NN threshold 約為 0.245，測試集 Recall 為 0.78，F2-score 為 0.430。也就是說，神經網路透過門檻調整後，才真正符合中風預警任務。",
    ),
    (
        "第 12 頁｜NN 混淆矩陣",
        "5:25-6:00",
        "從混淆矩陣可以更直接看到差異。預設 0.5 門檻下，TP 是 17、FN 是 33。調整到 0.245 後，TP 提升到 39、FN 降到 11，漏判減少 22 位。不過代價是 FP 從 53 增加到 215，代表系統會發出更多預警，後續需要醫療檢查流程來確認。",
    ),
    (
        "第 13 頁｜MLP 與 LR 比較",
        "6:00-6:35",
        "和 Logistic Regression 基準比較，MLP 的 PR-AUC 是 0.274，高於 LR 的 0.252，代表神經網路在風險排序上有進步。不過 LR 的 Recall 是 0.80、F2 是 0.445，略高於 MLP 的 Recall 0.78、F2 0.430。因此我的解讀是：神經網路是期末主軸，也確實帶來 PR-AUC 改善，但它不是無條件全面勝出，仍需要更深入調整。",
    ),
    (
        "第 14 頁｜期中基準與期末神經網路",
        "6:35-7:10",
        "再回到期中比較。期中的 Balanced LR 結果是 FP 251、FN 10、Recall 0.80。期末 MLP 是 FP 215、FN 11、Recall 0.78。也就是說，期末神經網路比期中少了 36 個誤報，但多漏判 1 位中風者。這不是推翻期中，而是延伸期中：從線性且好解釋的模型，進一步測試非線性模型是否有更好的風險排序能力。",
    ),
    (
        "第 15 頁｜NN 判讀",
        "7:10-7:40",
        "這頁是我對神經網路結果的判讀。MLP 的優點是能處理非線性特徵交互作用，PR-AUC 也較高；限制是中風樣本太少，神經網路容易受到類別不平衡影響，所以 Recall 和 F2 沒有明顯超過 LR。未來如果要讓 NN 更強，可以加入更多資料、做機率校準，或使用成本敏感訓練，讓漏判的代價在訓練時被更明確地考慮。",
    ),
    (
        "第 16 頁｜結論",
        "7:40-8:05",
        "最後總結。第一，MLP 適合這份表格型資料，可以學習非線性特徵組合。第二，經過 GridSearchCV 和 OOF F2 門檻調整後，神經網路測試集 Recall 達到 0.78，PR-AUC 達到 0.274。第三，NN 的 PR-AUC 高於 LR，但 Recall 和 F2 仍略低於 LR 基準。實務上，這個模型比較適合作為初步風險排序工具，提醒高風險者進一步檢查，而不能取代醫師診斷。",
    ),
    (
        "第 17 頁｜Thank you",
        "結尾",
        "以上是我的期末報告，謝謝大家。",
    ),
]


def set_run(run, size=11, bold=False, color=INK):
    run.font.name = "Microsoft JhengHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 24, NAVY, 0, 5),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 14, NAVY, 12, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("資料探勘期末報告｜8 分鐘講稿"), 9, False, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("王俊詠｜11127004"), 9, False, MUTED)

    p = doc.add_paragraph(style="Title")
    p.add_run("Stroke Prediction 神經網路延伸研究\n8 分鐘講稿")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("主軸：MLP Neural Network；比較基準：期中 Logistic Regression")

    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Inches(6.65)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(5)
    cp.paragraph_format.space_after = Pt(5)
    set_run(cp.add_run("報告策略：前半段交代資料與前處理，後半段集中說明 MLP、門檻調整、混淆矩陣與期中對比。"), 10.5, True, NAVY)

    for heading, timing, script in SLIDES:
        p = doc.add_paragraph(style="Heading 1")
        p.paragraph_format.keep_with_next = True
        p.add_run(heading)

        t = doc.add_paragraph()
        t.paragraph_format.space_after = Pt(2)
        set_run(t.add_run(timing), 10, True, RED)

        s = doc.add_paragraph()
        s.paragraph_format.line_spacing = 1.28
        s.paragraph_format.space_after = Pt(7)
        set_run(s.add_run(script), 11, False, INK)

    doc.add_page_break()
    doc.add_paragraph("上台提醒", style="Heading 1")
    checks = [
        "第 8 頁開始明確說：期末主軸是 MLP 神經網路。",
        "第 10-12 頁是重點：用 threshold=0.245 解釋 TP、FN 的改善。",
        "第 13-15 頁要誠實講：NN 的 PR-AUC 較高，但 Recall/F2 沒有贏過 LR。",
        "結論不要說 LR 是最終推薦主角；要說 NN 是期末主模型，LR 是比較基準。",
    ]
    for item in checks:
        p = doc.add_paragraph(style=None)
        set_run(p.add_run(f"- {item}"), 10.5, False, INK)

    doc.save(OUT)
    print(f"Created {OUT}")


if __name__ == "__main__":
    main()
