from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "11127004_王俊詠_Stroke_Prediction_逐頁註記與注意事項.docx"

NAVY = "20354A"
RED = "E4473D"
TEAL = "2A8C91"
MUTED = "687786"
PALE = "F2F4F7"


SLIDES = [
    {
        "page": 1,
        "title": "神經網路中風預測分析",
        "original": "無（封面特殊版型）。",
        "attention": "開場先說明這是從期中 Stroke Prediction 延伸而來；期末依課程要求加入神經網路，並以 MLPClassifier 作為主模型。Logistic Regression 僅作為 benchmark。",
        "numbers": "報告者：王俊詠；學號：11127004；報告時間約 8 分鐘。",
    },
    {
        "page": 2,
        "title": "期中建立基線，期末改以神經網路延伸",
        "original": "資料來源：期中報告與期末延伸研究。",
        "attention": "清楚切分兩階段：期中完成資料清理、視覺化與 LR／RF／KNN 基線；期末新增 MLP、GridSearchCV 與 OOF F2 threshold tuning。不要讓 LR 看起來像期末主模型。",
        "numbers": "期中：傳統 ML 基線；期末：MLP 主模型。",
    },
    {
        "page": 3,
        "title": "資料集的核心難題：中風案例只占 4.87%",
        "original": "資料來源：healthcare-dataset-stroke-data.csv。",
        "attention": "強調類別極度不平衡，因此 Accuracy 可能看起來很高，卻漏掉中風個案。後續應優先看 Recall、F2 與 PR-AUC。",
        "numbers": "總樣本 5,110；中風比例 4.87%；測試集 1,022。",
    },
    {
        "page": 4,
        "title": "視覺化顯示：風險集中在年齡與血糖等特徵",
        "original": "資料來源：本研究 EDA。",
        "attention": "可以說高齡與較高血糖的中風比例較明顯，但只能說資料中呈現關聯，不要宣稱因果關係。",
        "numbers": "重點特徵：age、avg_glucose_level，以及 hypertension、heart_disease 等風險資訊。",
    },
    {
        "page": 5,
        "title": "缺失值與相關性：先修補資料，再避免過度解讀",
        "original": "資料來源：本研究缺失值與相關性分析。",
        "attention": "BMI 缺失值必須在 Pipeline 內以訓練資料估計補值參數；相關係數只能描述線性關係，不能證明因果。",
        "numbers": "BMI 缺失值 201 筆；數值欄位使用中位數補值。",
    },
    {
        "page": 6,
        "title": "Pipeline 讓前處理只從訓練資料學習",
        "original": "方法：ColumnTransformer + Pipeline。",
        "attention": "說明類別欄位使用 One-hot encoding、數值欄位標準化，且補值與標準化都在每一折訓練資料內估計，以降低資料洩漏。",
        "numbers": "10 個原始特徵轉換為 21 個模型輸入特徵。",
    },
    {
        "page": 7,
        "title": "神經網路流程：測試集只在最後一次打開",
        "original": "方法：Stratified train/test split，random_state=42。",
        "attention": "模型架構、超參數與 threshold 都只能用訓練資料決定；測試集只用於最後一次泛化評估。這是避免資料洩漏的重要說明。",
        "numbers": "訓練集 4,088；測試集 1,022；分層切分。",
    },
    {
        "page": 8,
        "title": "神經網路模型設計：以 MLP 處理表格型中風資料",
        "original": "模型：sklearn MLPClassifier；資料：Healthcare Stroke Dataset。",
        "attention": "這是表格型資料，不是影像或序列，因此使用 MLP 比 CNN／RNN 更合理。MLP 是期末主模型，LR 僅是可解釋 benchmark。",
        "numbers": "主模型：Multilayer Perceptron（多層感知器）。",
    },
    {
        "page": 9,
        "title": "MLP GridSearchCV：搜尋神經網路架構與正則化",
        "original": "GridSearchCV scoring = average_precision；只使用訓練資料交叉驗證。",
        "attention": "說明 GridSearchCV 不是直接追求 Accuracy，而是使用 average precision／PR-AUC，以配合不平衡資料。",
        "numbers": "最佳設定：(64,) 隱藏層、ReLU、alpha=0.0001、learning_rate_init=0.001；CV PR-AUC=0.1639。",
    },
    {
        "page": 10,
        "title": "神經網路不能只看 0.5：中風預測需要召回導向門檻",
        "original": "門檻由訓練集 OOF 預測決定；測試集只用於最後評估。",
        "attention": "預設 threshold=0.5 並不一定適合醫療預警。F2-score 比 F1 更重視 Recall，因此用它選擇降低漏判的門檻。",
        "numbers": "預設門檻 FN=33；調整後 FN=11。",
    },
    {
        "page": 11,
        "title": "OOF F2 選出 NN threshold = 0.245",
        "original": "MLP Precision／Recall／F2-score threshold curve。",
        "attention": "指出紅色虛線 0.245 是訓練集 OOF F2-score 的最高附近，灰色虛線是預設 0.5。測試集沒有參與門檻選擇。",
        "numbers": "threshold=0.245；Recall=0.78；F2=0.430。",
    },
    {
        "page": 12,
        "title": "神經網路混淆矩陣：門檻調整後明顯降低漏判",
        "original": "測試集 n=1,022；MLP tuned threshold=0.245。",
        "attention": "先讀矩陣四格，再說明降低漏判的代價是誤警示增加。模型適合做初步風險預警，不能直接取代醫療診斷。",
        "numbers": "TN=757、FP=215、FN=11、TP=39；Recall=39/50=0.78。",
    },
    {
        "page": 13,
        "title": "結果分析：MLP 主模型與 LR benchmark 比較",
        "original": "比較基準：同一測試集、同一 train/test split；LR 僅作為 benchmark。",
        "attention": "MLP 的 PR-AUC 較高，但 Recall 與 F2 略低於 LR benchmark。只能說風險排序能力提升，不能說神經網路全面勝出。",
        "numbers": "MLP：PR-AUC=0.274、Recall=0.78、F2=0.430；LR benchmark：0.252、0.80、0.445。",
    },
    {
        "page": 14,
        "title": "期中基準 vs 期末神經網路：保留脈絡，但主軸轉向 MLP",
        "original": "期中數值來自 main.ipynb；期末 MLP 數值來自 neural_network_extension.ipynb。",
        "attention": "務必分清兩組 LR：期中 Balanced LR 的 FP=251；期末 LR benchmark 經 threshold tuning 後 FP=209。本頁右側則是期末 MLP 的 FP=215。",
        "numbers": "期中 Balanced LR：Recall=0.80、FN=10、FP=251；期末 MLP：Recall=0.78、FN=11、FP=215。",
    },
    {
        "page": 15,
        "title": "神經網路主軸下的判讀：MLP 有進步，也有取捨",
        "original": "模型：sklearn MLPClassifier；同一測試集比較。",
        "attention": "重申 MLP 是期末主模型；LR benchmark 與期中基線只是參考線。優勢是 PR-AUC，限制是 Precision 偏低且 False Positive 較多。",
        "numbers": "MLP PR-AUC=0.274、Recall=0.78、F2=0.430、FP=215。",
    },
    {
        "page": 16,
        "title": "結論：期末以神經網路延伸中風預測",
        "original": "無（結論特殊版型）。",
        "attention": "結論要保持誠實：MLP 在 PR-AUC 上高於 LR benchmark，但 Recall/F2 略低；因此不主張 NN 全面較強。未來可做外部驗證、機率校準與醫療成本導向門檻設定。",
        "numbers": "MLP PR-AUC=0.274；LR benchmark PR-AUC=0.252。",
    },
    {
        "page": 17,
        "title": "Thank you",
        "original": "無（結尾特殊版型）。",
        "attention": "簡短致謝後進入提問。預先準備回答：為何選 MLP、為何用 F2、為何 threshold 不是 0.5，以及為何 MLP 沒有全面勝過 LR。",
        "numbers": "姓名：王俊詠；學號：11127004。",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_font(run, name="Microsoft JhengHei", size=11, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_label_paragraph(cell, label, text, color):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label)
    set_font(r, size=10.5, bold=True, color=color)
    r2 = p.add_run(text)
    set_font(r2, size=10.5, color="1C2936")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft JhengHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 14, 7),
    ("Heading 2", 13, NAVY, 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Microsoft JhengHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(14)
title.paragraph_format.space_after = Pt(6)
r = title.add_run("Stroke Prediction 期末簡報\n逐頁註記與注意事項")
set_font(r, size=22, bold=True, color=NAVY)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(14)
r = meta.add_run("王俊詠｜11127004｜MLP Neural Network 主軸版")
set_font(r, size=11, bold=True, color=RED)

intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(12)
r = intro.add_run(
    "使用方式：這份文件收錄已從投影片左下角移除的原註記，並補充每頁口頭報告時需要注意的說法。"
    "投影片只保留右下角頁碼，方法與數據依據則集中在本文件備查。"
)
set_font(r, size=10.5, color=MUTED)

for item in SLIDES:
    heading = doc.add_paragraph(style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    r = heading.add_run(f"第 {item['page']} 頁｜{item['title']}")
    set_font(r, size=13, bold=True, color=NAVY)

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.18)
    table.columns[1].width = Inches(5.12)
    labels = [
        ("原頁尾註記", item["original"], PALE),
        ("報告時注意", item["attention"], "E8F2F2"),
        ("關鍵數字", item["numbers"], "FBE9E7"),
    ]
    for row, (label, value, fill) in zip(table.rows, labels):
        row.cells[0].width = Inches(1.18)
        row.cells[1].width = Inches(5.12)
        set_cell_shading(row.cells[0], fill)
        set_cell_shading(row.cells[1], "FFFFFF")
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(label)
        set_font(rr, size=9.5, bold=True, color=RED if label == "報告時注意" else NAVY)
        add_label_paragraph(row.cells[1], "", value, TEAL if label == "報告時注意" else MUTED)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for edge in ("top", "start", "bottom", "end"):
                tag = qn(f"w:{edge}")
                node = tc_mar.find(tag)
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "100")
                node.set(qn("w:type"), "dxa")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)

footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("11127004 王俊詠｜Stroke Prediction 期末簡報逐頁備註")
set_font(r, size=8.5, color=MUTED)

doc.save(OUTPUT)
print(OUTPUT)
