from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT = Path("11127004_王俊詠_Stroke_Prediction_延伸研究報告.docx")


def set_run_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Microsoft JhengHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def insert_paragraph_before(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p = new_p
    inserted._element = new_p
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def insert_table_before(paragraph, rows, cols):
    tbl = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addprevious(tbl._tbl)
    return tbl


def main():
    doc = Document(REPORT)
    marker = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("十、研究限制"):
            marker = p
            break
    if marker is None:
        raise RuntimeError("Cannot find insertion marker: 十、研究限制")

    # Avoid duplicated insertion if the script is run twice.
    if any("Neural Network 延伸比較" in p.text for p in doc.paragraphs):
        print("Report already contains Neural Network section.")
        return

    h = insert_paragraph_before(marker, "9.4 Neural Network 延伸比較", style="Heading 2")
    for run in h.runs:
        set_run_font(run, size=13, bold=True, color="2E74B5")

    paragraphs = [
        "為回應期末專案需加入神經網路的要求，本研究額外建立 MLPClassifier 作為延伸比較模型。由於 Stroke Prediction Dataset 屬於表格型資料，特徵包含年齡、平均血糖、BMI 與多個類別欄位，並非影像或時間序列，因此本研究選擇小型全連接神經網路（Multilayer Perceptron, MLP），而非 CNN、RNN 或 LSTM。",
        "MLP 使用與 Logistic Regression 相同的資料切分與前處理 Pipeline，包含中位數補值、One-Hot Encoding 與 StandardScaler。神經網路搜尋的最佳設定為 hidden_layer_sizes=(64,)、activation='relu'、alpha=0.0001、learning_rate_init=0.001，並使用 early stopping 降低過度擬合風險。由於中風類別樣本較少，訓練時同樣使用正類權重較高的 sample_weight，使模型更重視中風個案。",
        "實驗結果顯示，Neural Network 在測試集 PR-AUC 為 0.2736，略高於 Logistic Regression 的 0.2517；但在本研究最重視的篩檢指標上，MLP 的 Recall 為 0.78、F2-score 為 0.4295、False Negative 為 11，均略低於 Logistic Regression 調整門檻後的 Recall 0.80、F2-score 0.4454 與 False Negative 10。因此，神經網路可作為期末延伸比較模型，但最終推薦模型仍為 Logistic Regression + GridSearchCV + OOF threshold tuning。",
    ]
    for text in paragraphs:
        p = insert_paragraph_before(marker, text)
        for run in p.runs:
            set_run_font(run, size=11, color="000000")
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.25

    table = insert_table_before(marker, 4, 4)
    table.style = "Table Grid"
    data = [
        ["模型", "Recall", "F2-score", "False Negative"],
        ["Logistic Regression + threshold", "0.80", "0.4454", "10"],
        ["MLP Neural Network + threshold", "0.78", "0.4295", "11"],
        ["結論", "LR 較高", "LR 較高", "LR 較少漏判"],
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10.5, bold=(r == 0), color="000000")

    caption = insert_paragraph_before(marker, "表 2　Logistic Regression 與 MLP Neural Network 的測試集比較")
    for run in caption.runs:
        set_run_font(run, size=10, bold=True, color="687786")
    caption.paragraph_format.space_after = Pt(8)

    doc.save(REPORT)
    print(f"Updated {REPORT}")


if __name__ == "__main__":
    main()
