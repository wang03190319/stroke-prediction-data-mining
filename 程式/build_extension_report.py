from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("11127004_王俊詠_Stroke_Prediction_延伸研究報告.docx")
ASSETS = Path("extension_report_assets")

NAVY = "2F455C"
BLUE = "3E6B8F"
RED = "C51F1F"
LIGHT_BLUE = "E8EEF5"
LIGHT_RED = "FFF1EF"
LIGHT_GRAY = "F2F4F7"
GRAY = "697783"
BLACK = "202124"
WHITE = "FFFFFF"

FONT = "Microsoft JhengHei"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(
    doc,
    text,
    size=11,
    bold=False,
    color=BLACK,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before=0,
    after=7,
    line=1.25,
    italic=False,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    run = paragraph.add_run(text)
    set_run_font(run, size, bold, color, italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: NAVY, 2: BLUE, 3: NAVY}
    set_run_font(run, sizes[level], True, colors[level])
    return paragraph


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    label_run = paragraph.add_run(f"{label}　")
    set_run_font(label_run, 11, True, accent)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, 10.5, False, BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, filename, caption, width=6.3):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(ASSETS / filename), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(3)
    caption_p.paragraph_format.space_after = Pt(8)
    caption_p.paragraph_format.keep_with_next = True
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, 9.5, False, GRAY)


def add_metric_table(doc):
    rows = [
        ("Accuracy", "0.9198", "0.7857", "-0.1341"),
        ("Precision", "0.2647", "0.1606", "-0.1041"),
        ("Recall", "0.3600", "0.8000", "+0.4400"),
        ("F1-score", "0.3051", "0.2676", "-0.0375"),
        ("F2-score", "0.3358", "0.4454", "+0.1096"),
        ("PR-AUC", "0.2517", "0.2517", "不變"),
        ("False Negative", "32", "10", "-22"),
        ("False Positive", "50", "209", "+159"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [2400, 2400, 2400, 2160])
    headers = ["評估指標", "預設門檻 0.500", "調整門檻 0.268", "變化"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, 10, True, WHITE)
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        if row_index % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for column, text in enumerate(values):
            p = cells[column].paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            color = RED if column == 3 and text.startswith("+") else BLACK
            run = p.add_run(text)
            set_run_font(run, 10, column == 0, color)
    set_table_geometry(table, [2400, 2400, 2400, 2160])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_page_field(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, False, GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    suffix = paragraph.add_run(" 頁")
    set_run_font(suffix, 9, False, GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, NAVY, 14, 7),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("Stroke Prediction 延伸研究")
    set_run_font(run, 9, True, GRAY)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer_p)


def add_cover(doc):
    add_text(
        doc,
        "資料探勘期末延伸研究報告",
        size=12,
        bold=True,
        color=RED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=55,
        after=18,
    )
    add_text(
        doc,
        "中風預測模型之超參數\n與分類門檻最佳化",
        size=27,
        bold=True,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=12,
        line=1.15,
    )
    add_text(
        doc,
        "以 GridSearchCV 與 F2-score Threshold Tuning\n改善不平衡醫療資料的中風風險篩檢",
        size=14,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=42,
        line=1.25,
    )
    add_callout(
        doc,
        "核心成果",
        "最佳分類門檻約為 0.268；測試集 Recall 由 0.36 提升至 0.80，"
        "False Negative 由 32 人降至 10 人。",
        fill=LIGHT_RED,
        accent=RED,
    )
    add_text(
        doc,
        "報告人：王俊詠\n學號：11127004\n日期：2026 年 6 月",
        size=11,
        color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=65,
        after=8,
        line=1.5,
    )
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "摘要", 1)
    add_text(
        doc,
        "本研究延伸先前以 Stroke Prediction Dataset 進行的中風預測分析。"
        "期中研究發現，資料集中中風個案僅占 4.87%，原始分類模型容易偏向未中風類別；"
        "雖然使用 class_weight 或 SMOTE 可提高 Recall，但模型仍採用固定參數與預設分類門檻 0.5。"
        "因此，本研究進一步使用 GridSearchCV 搜尋 Logistic Regression 的 C、class_weight 與 solver，"
        "並以五折交叉驗證 PR-AUC 作為模型選擇指標。完成模型選擇後，再利用訓練集的 Out-of-Fold"
        " 預測機率尋找 F2-score 最高的分類門檻，避免使用測試集調整參數造成資料洩漏。",
    )
    add_text(
        doc,
        "實驗結果顯示，最佳參數為 C=0.01、class_weight={0:1, 1:5}、solver=liblinear，"
        "交叉驗證 PR-AUC 為 0.1958；最佳分類門檻約為 0.268。相較於預設門檻 0.5，"
        "調整門檻後測試集 Recall 由 0.36 提升至 0.80，F2-score 由 0.34 提升至 0.45，"
        "False Negative 由 32 人下降至 10 人。然而，Precision 由 0.26 降至 0.16，"
        "False Positive 由 50 人增加至 209 人。結果說明，分類門檻調整可以顯著降低漏判，"
        "但必須接受更多誤警示。本模型應定位為中風風險的輔助篩檢工具，而非臨床診斷工具。",
    )
    add_callout(
        doc,
        "關鍵字",
        "中風預測、類別不平衡、Logistic Regression、GridSearchCV、PR-AUC、"
        "分類門檻、F2-score、Out-of-Fold",
    )

    add_heading(doc, "一、研究背景與延伸動機", 1)
    add_text(
        doc,
        "中風屬於高風險醫療事件，若模型將真正有中風風險的個案判定為未中風，"
        "可能延誤後續檢查與治療。期中研究使用 5,110 筆資料進行分析，其中未中風者有"
        " 4,861 人，中風者僅 249 人，形成明顯的類別不平衡。這使 Accuracy 容易產生"
        "「看似很高但缺乏篩檢價值」的現象，因此醫療篩檢應特別重視 Recall 與 False Negative。",
    )
    add_text(
        doc,
        "期中報告已比較 Logistic Regression、Random Forest 與 KNN，並使用 class_weight"
        " 與 SMOTE 改善少數類別辨識能力。其研究限制指出，尚未進行完整的超參數搜尋，"
        "且所有模型仍以 0.5 作為固定分類門檻。本次延伸研究即針對這兩項限制提出改善。",
    )
    add_callout(
        doc,
        "與期中報告的銜接",
        "期中流程使用 BMI 平均值填補、pandas One-Hot Encoding、MinMaxScaler，並比較多種模型"
        "與 SMOTE；本次延伸改用可避免交叉驗證洩漏的 Pipeline，以中位數填補、"
        "OneHotEncoder 與 StandardScaler 處理資料，再針對期中推薦的 Logistic Regression"
        "進行 GridSearchCV 與 OOF 門檻最佳化。",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )
    add_heading(doc, "二、研究目的", 1)
    add_text(
        doc,
        "本研究的第一個目的，是透過 GridSearchCV 系統化比較 Logistic Regression 的"
        "超參數組合，而不是以人工方式選擇單一設定。第二個目的，是在不使用測試集資訊的"
        "前提下調整分類門檻，使模型更符合中風早期風險篩檢對 Recall 的需求。第三個目的，"
        "是量化提高 Recall 所伴隨的 Precision、Accuracy 與 False Positive 代價。",
    )
    add_text(
        doc,
        "本次研究同時升級前處理方式、模型參數與分類門檻，因此期中與期末結果屬於完整流程的"
        "延伸比較，不是只改變單一變因的消融實驗。期末報告中「門檻調整前後」的比較，"
        "則固定使用同一個 GridSearchCV 最佳模型，只改變 0.5 與 0.268 兩個決策門檻，"
        "可直接反映 threshold tuning 的影響。",
    )
    doc.add_page_break()

    add_heading(doc, "三、Dataset 與資料視覺化", 1)
    add_heading(doc, "3.1 Dataset 介紹", 2)
    add_text(
        doc,
        "研究使用 healthcare-dataset-stroke-data.csv，共 5,110 筆資料。id 欄位不具預測意義，"
        "因此自特徵中移除；stroke 為二元目標變數。模型特徵包含人口資料、健康狀況與生活型態，"
        "例如 age、gender、hypertension、heart_disease、avg_glucose_level、bmi、"
        "work_type 與 smoking_status。資料中未中風者有 4,861 人，中風者僅 249 人，"
        "中風比例為 4.87%，屬於明顯的類別不平衡分類問題。",
    )
    add_heading(doc, "3.2 視覺化分析", 2)
    add_figure(
        doc,
        "00_dataset_eda.png",
        "圖 1　資料類別分布、年齡與平均血糖值的中風分組比較",
        6.25,
    )
    add_text(
        doc,
        "類別分布圖顯示，中風資料遠少於未中風資料，若模型只追求 Accuracy，"
        "可能傾向將所有個案預測為未中風。年齡箱型圖顯示，中風族群的年齡中位數明顯較高；"
        "平均血糖圖也顯示中風族群的分布範圍較廣且上四分位數較高。這些結果指出，"
        "年齡與代謝健康可能是重要風險訊號。",
    )
    doc.add_page_break()

    add_heading(doc, "四、資料前處理與相關性分析", 1)
    add_heading(doc, "4.1 Missing Value", 2)
    add_text(
        doc,
        "資料中只有 bmi 欄位存在缺失值，共 201 筆。為保留少數中風樣本，本研究不直接刪除缺失資料，"
        "而是在 Pipeline 中使用 SimpleImputer(strategy='median') 以訓練資料中位數填補。"
        "相較平均值，中位數較不容易受到 BMI 極端值影響。由於填補程序位於 Pipeline 內，"
        "交叉驗證每一折都只使用該折訓練資料估計中位數。",
    )
    add_heading(doc, "4.2 相關性分析", 2)
    add_figure(
        doc,
        "00_correlation_heatmap.png",
        "圖 2　主要數值特徵與 stroke 的相關性熱圖",
        5.55,
    )
    add_text(
        doc,
        "stroke 與 age 的線性相關係數約為 0.25，為數值特徵中最高；"
        "hypertension、heart_disease 與 avg_glucose_level 均約為 0.13。"
        "BMI 與 stroke 的線性相關較低，約為 0.04。相關性只能說明線性關聯，"
        "不能直接解釋因果關係，也不能單獨代表特徵在模型中的完整預測價值。",
    )
    add_heading(doc, "4.3 One-Hot Encoding 與 Standardization", 2)
    add_figure(
        doc,
        "00_preprocessing_summary.png",
        "圖 3　缺失值處理、One-Hot Encoding 與 StandardScaler 摘要",
        6.25,
    )
    add_text(
        doc,
        "類別欄位使用 OneHotEncoder(handle_unknown='ignore') 轉為模型可處理的數值特徵，"
        "原始 10 個模型欄位經轉換後形成 21 個特徵。數值欄位使用 StandardScaler，"
        "將各欄位轉換為平均數接近 0、標準差接近 1 的尺度。本研究採用的是標準化"
        "（standardization），不是將資料壓縮到 0 到 1 的 Min-Max normalization。"
        "統一尺度可避免 age、avg_glucose_level 與 bmi 的數值範圍差異影響 Logistic Regression。",
    )
    add_text(
        doc,
        "此處與期中流程不同：期中使用全資料平均值填補 BMI，並在切分前以 MinMaxScaler"
        "進行 0 到 1 縮放；延伸程式則將補值、編碼與標準化放入 Pipeline，讓每一折交叉驗證"
        "只從該折訓練資料學習前處理參數。這是本次研究在方法嚴謹度上的另一項延伸。",
    )
    doc.add_page_break()

    add_heading(doc, "五、ML 分類模型與研究方法", 1)
    add_heading(doc, "5.1 分類模型選擇", 2)
    add_text(
        doc,
        "本研究屬於分類問題，目標為預測 stroke=0 或 stroke=1。期中研究已比較"
        " Logistic Regression、Random Forest 與 KNN，結果顯示 Logistic Regression"
        " 搭配不平衡權重能取得較高 Recall，且模型結構相對容易解釋。因此，期末延伸研究"
        "選擇 Logistic Regression 作為主要模型，集中分析參數調整與分類門檻的影響。",
    )
    add_text(
        doc,
        "期末延伸模型沒有再使用 SMOTE 產生合成樣本，而是將 class_weight 納入 GridSearchCV。"
        "此設計延續期中報告推薦 class_weight 方法的結論，同時讓模型從交叉驗證中選擇"
        " None、balanced、1:3、1:5 或 1:10 的權重設定。",
    )
    add_heading(doc, "5.2 資料切分方式", 2)
    add_text(
        doc,
        "資料以 80% 作為訓練集、20% 作為測試集，並使用 stratify=y 維持中風比例。"
        "訓練集共有 4,088 筆，測試集共有 1,022 筆，其中測試集包含 50 位中風者。"
        "測試集在完成所有模型與門檻選擇前保持未使用狀態。",
    )
    add_heading(doc, "5.3 前處理 Pipeline", 2)
    add_text(
        doc,
        "數值特徵包含 age、hypertension、heart_disease、avg_glucose_level 與 bmi。"
        "數值缺失值以中位數填補，再使用 StandardScaler 標準化。類別特徵包含 gender、"
        "ever_married、work_type、Residence_type 與 smoking_status；類別缺失值以眾數填補，"
        "再使用 One-Hot Encoding。所有前處理與 Logistic Regression 結合為 Pipeline，"
        "確保每一折交叉驗證都只使用該折訓練資料估計補值與標準化參數，降低資料洩漏風險。",
    )
    add_heading(doc, "5.4 GridSearchCV 設計", 2)
    add_text(
        doc,
        "本研究搜尋五種 C 值（0.01、0.1、1、10、100）、五種 class_weight 設定"
        "（None、balanced、1:3、1:5、1:10）與兩種 solver（liblinear、lbfgs），"
        "共 50 組參數。每組參數使用五折 Stratified Cross-Validation 評估，"
        "並以 average precision，即 PR-AUC，作為最佳模型選擇標準。PR-AUC 特別關注少數類別"
        "的 Precision 與 Recall，較 Accuracy 適合本研究的不平衡資料。",
    )
    add_figure(
        doc,
        "01_gridsearch_pr_auc_heatmap.png",
        "圖 4　GridSearchCV 各參數組合的五折交叉驗證 PR-AUC",
        6.2,
    )
    doc.add_page_break()

    add_heading(doc, "六、參數調整結果", 1)
    add_text(
        doc,
        "GridSearchCV 的最佳參數組合為 C=0.01、class_weight={0:1, 1:5}、"
        "solver=liblinear，平均交叉驗證 PR-AUC 為 0.1958。C=0.01 代表較強的正則化，"
        "理論上可降低模型複雜度並抑制過度擬合；將中風類別權重提高為 5 倍，"
        "則使模型在訓練時更重視中風個案的分類錯誤。",
    )
    add_text(
        doc,
        "熱圖也顯示，多數參數組合的 PR-AUC 約介於 0.19 至 0.196，差異並不大。"
        "因此，本研究不應將最佳參數解讀為大幅提升模型辨識能力，而應視為在既有資料與"
        "Logistic Regression 架構下，找到相對穩定且較符合少數類別需求的設定。",
    )
    add_callout(
        doc,
        "方法意義",
        "超參數搜尋改善的是模型本身的設定；分類門檻調整則改變預測機率轉換為類別的決策規則。"
        "兩者必須分開處理，才能避免將模型訓練與實際應用決策混為一談。",
        fill=LIGHT_RED,
        accent=RED,
    )

    add_heading(doc, "七、分類門檻最佳化", 1)
    add_heading(doc, "7.1 為何不直接使用 0.5？", 2)
    add_text(
        doc,
        "Logistic Regression 輸出的是中風機率，預設門檻 0.5 只是一項一般性設定，"
        "並不代表最適合醫療篩檢。當中風樣本稀少時，許多真正中風個案的預測機率可能低於 0.5，"
        "造成大量漏判。降低門檻會讓模型更容易判定為中風，通常能提高 Recall，"
        "但也會增加 False Positive。",
    )
    add_heading(doc, "7.2 OOF 與 F2-score", 2)
    add_text(
        doc,
        "為避免使用測試集選擇門檻，本研究在訓練集上再次執行五折交叉驗證，"
        "產生 Out-of-Fold 預測機率。每一筆訓練資料的機率都由未看過該筆資料的模型產生，"
        "因此可較合理地模擬未知資料。接著計算各門檻下的 Precision、Recall 與 F2-score。"
        "F2-score 對 Recall 的重視程度高於 Precision，符合本研究降低漏判的目標。",
    )
    add_figure(
        doc,
        "02_threshold_precision_recall_f2.png",
        "圖 5　不同分類門檻下的 Precision、Recall 與 F2-score",
        6.2,
    )
    add_text(
        doc,
        "如圖 5 所示，最佳 F2-score 對應的門檻約為 0.268。在該 OOF 門檻位置，"
        "Recall 約為 0.769、Precision 約為 0.145。這個結果反映研究的決策偏好："
        "寧可增加部分誤警示，也希望降低真正中風個案被漏掉的機率。",
    )
    doc.add_page_break()

    add_heading(doc, "八、資料結果分析", 1)
    add_heading(doc, "8.1 混淆矩陣比較", 2)
    add_figure(
        doc,
        "03_confusion_matrix_comparison.png",
        "圖 6　預設門檻與調整後門檻的測試集混淆矩陣",
        6.15,
    )
    add_text(
        doc,
        "在預設門檻 0.5 下，模型正確找出 18 位中風者，但漏掉 32 位中風者，"
        "中風類別 Recall 為 0.36。改用 0.268 後，模型成功找出 40 位中風者，"
        "僅漏掉 10 位，Recall 提升至 0.80。換言之，門檻調整多找出 22 位真正中風者，"
        "同時減少 22 位漏判。",
    )
    add_text(
        doc,
        "然而，未中風者被誤判為中風的人數由 50 人增加至 209 人。這表示降低門檻"
        "提高了模型敏感度，但也使醫療系統可能需要處理更多後續檢查與警示。此結果不能只以"
        "Recall 提升判定為全面改善，而必須同時評估誤警示成本。",
    )

    add_heading(doc, "8.2 Precision、Recall、F1 與 F2 比較", 2)
    add_metric_table(doc)
    add_text(
        doc,
        "分類問題依老師規定以 Precision、F1-score 與混淆矩陣為主要結果，"
        "並額外呈現 Recall、F2-score、Accuracy 與 PR-AUC。調整門檻後 F1-score"
        "由 0.3051 降至 0.2676，原因是 Precision 的下降幅度大於 Recall 的提升對 F1 的貢獻；"
        "但更重視 Recall 的 F2-score 則由 0.3358 提升至 0.4454。"
        "PR-AUC 在兩個門檻下皆為 0.2517，因為 PR-AUC 是由完整預測機率排序計算，"
        "不會因單一分類門檻而改變。Accuracy、Precision、Recall 與 F2-score 則依賴最終類別，"
        "因此會隨門檻調整。這也說明評估模型時，必須區分「機率排序能力」與「特定門檻下的分類結果」。",
    )
    doc.add_page_break()

    add_heading(doc, "九、醫療篩檢取捨與討論", 1)
    add_figure(
        doc,
        "04_performance_tradeoff.png",
        "圖 7　分類門檻調整前後的核心指標與誤判取捨",
        6.2,
    )
    add_heading(doc, "9.1 提高 Recall 的價值", 2)
    add_text(
        doc,
        "在中風風險篩檢情境中，False Negative 代表真正有中風風險的個案沒有被模型辨識。"
        "本研究將 FN 由 32 人降低至 10 人，顯示門檻調整能有效改善漏判問題。"
        "如果模型用於第一階段風險篩選，而後續仍有醫師判讀或檢查，較高 Recall 具有實際價值。",
    )
    add_heading(doc, "9.2 誤警示與 Precision 的代價", 2)
    add_text(
        doc,
        "門檻調整後 Precision 降至 0.1606，表示模型判定為中風風險者中，"
        "真正中風者約占 16%。False Positive 增加至 209 人，可能造成額外檢查、"
        "醫療資源負擔與個案焦慮。因此，0.268 並非唯一正確門檻，而是基於 F2-score 與"
        "「降低漏判優先」假設得到的選擇。實際部署時應由醫療成本與風險政策共同決定門檻。",
    )
    add_heading(doc, "9.3 與期中推薦模型的關係", 2)
    add_text(
        doc,
        "期中報告推薦 Logistic Regression + class_weight='balanced'，其測試結果為"
        " TP=40、FN=10、FP=251，Recall 同樣為 0.80。本次延伸流程得到 TP=40、FN=10、"
        "FP=209，在維持 Recall 0.80 的同時減少 42 個 False Positive，Precision 也由期中約"
        " 0.14 提升至 0.16。不過，由於前處理與參數設定同時改變，這項差異應解讀為完整延伸流程"
        "的結果，不能完全歸因於單一超參數或分類門檻。",
    )
    add_callout(
        doc,
        "應用定位",
        "本模型適合作為風險提示或二次檢查的觸發依據，不應單獨用來診斷中風。"
        "模型輸出必須搭配醫療專業判斷與後續檢查。",
        fill=LIGHT_RED,
        accent=RED,
    )

    add_heading(doc, "十、研究限制", 1)
    add_text(
        doc,
        "第一，資料集中只有 249 位中風者，少數類別樣本有限，交叉驗證結果可能受到資料切分影響。"
        "第二，本研究只延伸 Logistic Regression，尚未對 XGBoost、LightGBM 或其他進階模型"
        "進行同等程度的參數與門檻比較。第三，最佳門檻只在單一資料集上取得，缺乏外部資料驗證，"
        "不能直接推論到其他醫院或族群。第四，F2-score 對 Recall 的偏好是研究設定，"
        "尚未納入真實醫療成本、資源容量與病患風險。第五，資料缺少血壓實際數值、膽固醇、"
        "家族病史與用藥紀錄等臨床特徵，限制模型預測能力。",
    )
    doc.add_page_break()

    add_heading(doc, "十一、結論：隱藏趨勢與未來應用", 1)
    add_text(
        doc,
        "本研究完成期中報告所提出的兩項延伸方向：超參數最佳化與分類門檻調整。"
        "GridSearchCV 找到的最佳 Logistic Regression 參數為 C=0.01、"
        "class_weight={0:1, 1:5}、solver=liblinear；接著利用訓練集 OOF 預測與 F2-score"
        " 選出約 0.268 的分類門檻。",
    )
    add_text(
        doc,
        "在未參與調參的測試集上，調整門檻使 Recall 由 0.36 提升至 0.80，"
        "F2-score 由 0.34 提升至 0.45，False Negative 由 32 人下降至 10 人。"
        "此結果證明，在嚴重類別不平衡的醫療資料中，模型訓練完成後的決策門檻同樣重要。"
        "但 Precision 下降與 False Positive 增加也顯示，高敏感度並非沒有成本。",
    )
    add_text(
        doc,
        "整體而言，若目標是早期風險篩檢，0.268 門檻比預設 0.5 更符合降低漏判的需求；"
        "若目標是減少不必要警示，則需要選擇更高門檻。未來研究應加入外部驗證、更多臨床特徵、"
        "成本敏感分析與模型校準，並由實際醫療情境決定可接受的 Recall 與 Precision 平衡。",
    )
    add_text(
        doc,
        "本研究觀察到的隱藏趨勢包括：中風風險與年齡的關聯最明顯，高血壓、心臟病與平均血糖"
        "亦呈正相關；同時，模型在類別不平衡資料下的問題不只來自演算法，也來自預設決策門檻。"
        "未來可將模型應用於健檢或門診的初步風險分層，將高風險個案交由醫療人員進一步評估，"
        "但不可取代正式診斷。",
    )
    add_callout(
        doc,
        "最終結論",
        "超參數搜尋決定模型如何學習；分類門檻決定模型如何行動。"
        "對不平衡醫療資料而言，兩者缺一不可。",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )

    add_heading(doc, "十二、未來研究方向", 1)
    future_rows = [
        ("外部驗證", "使用不同醫院、地區或年份的資料檢驗模型泛化能力。"),
        ("模型比較", "將相同 Grid Search 與 threshold 流程套用至 XGBoost、LightGBM 與 SVM。"),
        ("機率校準", "使用 Platt scaling 或 isotonic regression 改善風險機率可信度。"),
        ("成本敏感決策", "以漏判與誤警示的實際成本建立門檻選擇函數。"),
        ("臨床特徵擴充", "加入血壓、膽固醇、家族史、用藥與生活習慣等資訊。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 6860])
    for index, text in enumerate(["方向", "內容"]):
        set_cell_shading(table.rows[0].cells[index], NAVY)
        p = table.rows[0].cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, 10, True, WHITE)
    for row_index, (label, detail) in enumerate(future_rows, start=1):
        cells = table.add_row().cells
        if row_index % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for index, text in enumerate((label, detail)):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, 10, index == 0, BLACK)
    set_table_geometry(table, [2500, 6860])

    add_heading(doc, "附錄：實驗重現資訊", 1)
    add_text(
        doc,
        "主要分析 Notebook：gridsearch_threshold.ipynb\n"
        "製圖程式：generate_extension_charts.py\n"
        "資料檔：healthcare-dataset-stroke-data.csv\n"
        "隨機種子：42\n"
        "訓練／測試比例：80%／20%\n"
        "交叉驗證：五折 StratifiedKFold\n"
        "模型選擇指標：average precision（PR-AUC）\n"
        "門檻選擇指標：F2-score",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line=1.4,
    )

    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_report()
