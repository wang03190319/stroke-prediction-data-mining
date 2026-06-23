from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def report_path() -> Path:
    candidates = [
        p
        for p in Path(".").glob("*.docx")
        if p.name.startswith("11127004_") and "Stroke_Prediction" in p.name and p.stat().st_size > 100_000
    ]
    if not candidates:
        raise FileNotFoundError("Cannot find final report docx.")
    return max(candidates, key=lambda p: p.stat().st_size)


def set_paragraph_text(paragraph, text: str):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    run = paragraph.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    return run


def main():
    path = report_path()
    doc = Document(path)
    p = doc.paragraphs

    replacements = {
        1: "神經網路中風預測分析\n以 MLP 作為期末延伸主軸",
        2: "以 GridSearchCV、OOF F2 Threshold Tuning\n評估不平衡醫療資料的神經網路風險篩檢",
        7: (
            "本研究延伸先前以 Stroke Prediction Dataset 進行的中風預測分析。期中研究發現，資料集中中風個案僅占 4.87%，"
            "模型容易偏向未中風類別，因此不能只依賴 Accuracy。期末研究依老師要求加入神經網路，並將 MLPClassifier "
            "設定為主要模型；Logistic Regression 則保留為期中基準與可解釋比較模型。"
        ),
        8: (
            "神經網路最佳設定為 hidden_layer_sizes=(64,)、activation='relu'、alpha=0.0001、learning_rate_init=0.001；"
            "OOF F2 選出的分類門檻約為 0.245。測試集結果顯示，MLP 的 Recall 為 0.78、F2-score 為 0.4295、"
            "PR-AUC 為 0.2736；相較 Logistic Regression 基準，MLP 的 PR-AUC 較高，但 Recall 與 F2 略低。"
            "因此，本研究重點在於以神經網路延伸期中分析，並誠實呈現非線性模型在不平衡醫療資料中的優勢與限制。"
        ),
        15: (
            "本研究的第一個目的，是在期中分析基礎上建立 MLP Neural Network，檢驗非線性模型是否能改善中風風險排序。"
            "第二個目的，是透過 GridSearchCV 系統化搜尋 MLP 架構、activation、regularization 與 learning rate。"
            "第三個目的，是在不使用測試集資訊的前提下，以 OOF F2-score 調整分類門檻，使模型更符合中風早期風險篩檢對 Recall 的需求。"
        ),
        16: (
            "本次研究同時保留 Logistic Regression 作為 benchmark，原因是期中模型具有可解釋性，也能提供神經網路是否真正改善的參考線。"
            "因此，期末結果應解讀為「以 MLP 為主模型、LR 為比較基準」的延伸研究，而不是只替換單一參數的消融實驗。"
        ),
        41: (
            "本研究屬於分類問題，目標為預測 stroke=0 或 stroke=1。期中研究已比較 Logistic Regression、Random Forest 與 KNN；"
            "期末研究則依要求加入 Neural Network，並選擇 MLPClassifier 作為主要模型。由於資料屬於表格型欄位資料，"
            "並非影像或時間序列，因此選擇全連接神經網路 MLP，而不是 CNN 或 RNN。"
        ),
        42: (
            "MLP 可學習年齡、平均血糖、BMI、工作型態與吸菸狀態等欄位之間的非線性組合。"
            "同時，為了讓期末結果能與期中銜接，本研究保留 Logistic Regression + threshold tuning 作為比較基準。"
        ),
        47: "5.4 Neural Network GridSearchCV 設計",
        48: (
            "MLP GridSearchCV 搜尋 hidden_layer_sizes、activation、alpha 與 learning_rate_init，"
            "並使用三折 Stratified Cross-Validation 評估。模型選擇指標使用 average precision，也就是 PR-AUC，"
            "因為它比 Accuracy 更能反映少數中風類別的排序能力。Logistic Regression 的 GridSearchCV 結果則作為 benchmark 保留。"
        ),
        52: "六、神經網路參數調整結果",
        53: (
            "MLP 的最佳參數組合為 hidden_layer_sizes=(64,)、activation='relu'、alpha=0.0001、"
            "learning_rate_init=0.001，交叉驗證 PR-AUC 為 0.1639。此設定代表模型使用單層 64 個神經元，"
            "以 ReLU 捕捉非線性關係，並透過 L2 regularization 降低過度擬合風險。"
        ),
        54: (
            "由於中風樣本只有 249 筆，神經網路不應被期待在所有指標上一定勝過線性模型。"
            "本研究將 MLP 視為期末主模型，重點在於觀察非線性模型能否提升風險排序能力，"
            "並透過 LR benchmark 檢查 Recall、F2 與 False Negative 是否仍符合醫療篩檢需求。"
        ),
        58: (
            "MLP 輸出的是中風風險機率，預設門檻 0.5 只是一項一般性設定，並不代表最適合醫療篩檢。"
            "在測試集上，若直接使用 0.5，神經網路只抓到 17 位中風者，False Negative 達 33。"
            "因此本研究使用 OOF F2-score 選擇 threshold=0.245，以降低漏判風險。"
        ),
        63: (
            "最佳 MLP 門檻約為 0.245。這個結果反映研究的決策偏好：在中風預警任務中，"
            "寧可增加部分 False Positive，也希望減少真正中風個案被漏掉的機率。"
            "門檻選擇只使用訓練集 OOF 預測，測試集保留到最後評估。"
        ),
        69: (
            "在 MLP 預設門檻 0.5 下，模型正確找出 17 位中風者，但漏掉 33 位中風者。"
            "改用 OOF F2 選出的 0.245 後，模型成功找出 39 位中風者，僅漏掉 11 位，Recall 達 0.78。"
            "換言之，門檻調整讓神經網路多找出 22 位真正中風者，同時減少 22 位漏判。"
        ),
        70: (
            "然而，未中風者被誤判為中風的人數由 53 人增加至 215 人。這表示降低門檻提高了模型敏感度，"
            "但也使醫療系統可能需要處理更多後續檢查與警示。此結果不能只以 Recall 提升判定為全面改善，"
            "而必須同時評估誤警示成本。"
        ),
        79: (
            "在中風風險篩檢情境中，False Negative 代表真正有中風風險的個案沒有被模型辨識。"
            "本研究將 MLP 的 FN 由 33 人降低至 11 人，顯示 threshold tuning 能有效改善神經網路的漏判問題。"
            "如果模型用於第一階段風險篩選，而後續仍有醫師判讀或檢查，較高 Recall 具有實際價值。"
        ),
        80: (
            "MLP 門檻調整後 Precision 為 0.1535，False Positive 增加至 215 人，可能造成額外檢查、醫療資源負擔與個案焦慮。"
            "因此，0.245 並非唯一正確門檻，而是基於 F2-score 與「降低漏判優先」假設得到的選擇。"
            "實際部署時應由醫療成本與風險政策共同決定門檻。"
        ),
        82: "9.3 與期中推薦模型的關係",
        83: (
            "期中報告推薦 Balanced Logistic Regression，其測試結果為 TP=40、FN=10、FP=251，Recall 為 0.80。"
            "期末 MLP 的測試結果為 TP=39、FN=11、FP=215，Recall 為 0.78。"
            "與期中相比，MLP 減少 36 個 False Positive，但多 1 個 False Negative；同時，MLP 的 PR-AUC 0.2736 高於 LR benchmark 的 0.2517。"
            "因此，期末神經網路的價值主要在於風險排序能力提升，而不是在所有分類指標上完全勝出。"
        ),
        85: "9.4 Neural Network 主模型比較",
        86: (
            "為回應期末專案需加入神經網路的要求，本研究將 MLPClassifier 作為期末主模型。"
            "由於 Stroke Prediction Dataset 屬於表格型資料，特徵包含年齡、平均血糖、BMI 與多個類別欄位，"
            "並非影像或時間序列，因此本研究選擇小型全連接神經網路（Multilayer Perceptron, MLP），而非 CNN、RNN 或 LSTM。"
        ),
        88: (
            "實驗結果顯示，Neural Network 在測試集 PR-AUC 為 0.2736，高於 Logistic Regression 的 0.2517；"
            "但在篩檢指標上，MLP 的 Recall 為 0.78、F2-score 為 0.4295、False Negative 為 11，"
            "略低於 Logistic Regression 調整門檻後的 Recall 0.80、F2-score 0.4454 與 False Negative 10。"
            "因此，本研究以 MLP 作為期末重點模型，同時以 LR 作為 benchmark，呈現神經網路在風險排序上改善、但仍受類別不平衡限制的結果。"
        ),
        91: (
            "第一，資料集中只有 249 位中風者，少數類別樣本有限，神經網路結果可能受到資料切分影響。"
            "第二，本研究使用小型 MLP，尚未進一步比較更深層網路、機率校準或成本敏感 loss。"
            "第三，最佳門檻只在單一資料集上取得，缺乏外部資料驗證，不能直接推論到其他醫院或族群。"
            "第四，F2-score 對 Recall 的偏好是研究假設，實際部署仍需醫療成本與風險政策共同決定。"
        ),
        94: (
            "本研究完成期中報告的延伸，並依期末要求將重點轉向神經網路。"
            "MLP GridSearchCV 找到的最佳參數為 hidden_layer_sizes=(64,)、activation='relu'、alpha=0.0001、learning_rate_init=0.001；"
            "接著利用訓練集 OOF 預測與 F2-score 選出約 0.245 的分類門檻。"
        ),
        95: (
            "在未參與調參的測試集上，MLP threshold tuning 使 False Negative 由 33 人下降至 11 人，"
            "True Positive 由 17 人提升至 39 人，Recall 達 0.78，F2-score 達 0.4295。"
            "同時，MLP 的 PR-AUC 為 0.2736，高於 LR benchmark 的 0.2517，顯示神經網路在風險排序上具有延伸價值。"
        ),
        96: (
            "整體而言，若目標是早期風險篩檢，MLP 搭配 0.245 門檻可作為初步風險排序工具；"
            "若目標是減少不必要警示，則需要選擇更高門檻或加入機率校準。"
            "未來研究應加入外部驗證、更多臨床特徵、成本敏感分析與模型校準，並由實際醫療情境決定可接受的 Recall 與 Precision 平衡。"
        ),
        97: (
            "本研究觀察到的隱藏趨勢包括：中風風險與年齡的關聯最明顯，高血壓、心臟病與平均血糖亦呈正相關；"
            "同時，模型在類別不平衡資料下的問題不只來自演算法，也來自預設決策門檻。"
            "神經網路能捕捉非線性特徵組合，但仍需門檻調整與醫療情境判讀。未來可將模型應用於健檢或門診的初步風險分層，"
            "將高風險個案交由醫療人員進一步評估，但不可取代正式診斷。"
        ),
        101: (
            "主要分析 Notebook：neural_network_extension.ipynb、gridsearch_threshold.ipynb\n"
            "神經網路執行 Notebook：neural_network_extension_executed.ipynb\n"
            "製圖程式：generate_extension_charts.py、neural_network_extension.py\n"
            "資料檔：healthcare-dataset-stroke-data.csv\n"
            "隨機種子：42\n"
            "訓練／測試比例：80%／20%\n"
            "神經網路交叉驗證：三折 StratifiedKFold\n"
            "模型選擇指標：average precision / PR-AUC\n"
            "門檻選擇：Out-of-Fold F2-score"
        ),
    }

    for idx, text in replacements.items():
        run = set_paragraph_text(p[idx], text)
        if idx in {1, 2}:
            run.bold = True
        if idx in {1}:
            run.font.size = Pt(24)
        elif idx in {2}:
            run.font.size = Pt(13)

    try:
        doc.save(path)
        print(f"Updated {path}")
    except PermissionError:
        fallback = path.with_name(path.stem + "_NN主軸版.docx")
        doc.save(fallback)
        print(f"Original report is locked; saved updated copy to {fallback}")


if __name__ == "__main__":
    main()
