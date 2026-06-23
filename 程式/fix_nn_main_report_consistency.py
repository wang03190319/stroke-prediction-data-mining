from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


REPORT_HINT = "NN主軸版"


def find_report() -> Path:
    candidates = [
        p
        for p in Path(".").glob("*.docx")
        if p.name.startswith("11127004_")
        and "Stroke_Prediction" in p.name
        and p.stat().st_size > 100_000
    ]
    nn = [p for p in candidates if REPORT_HINT in p.name]
    return max(nn or candidates, key=lambda p: p.stat().st_mtime)


def set_run_style(run, size=10.5, bold=False):
    run.font.name = "Microsoft JhengHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph(paragraph, text: str, size=11, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_style(run, size=size, bold=bold)
    return paragraph


def set_cell(cell, text: str, size=10.5, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    set_run_style(run, size=size, bold=bold)


def normalize_table(table, rows: int, cols: int):
    while len(table.rows) < rows:
        table.add_row()
    # python-docx cannot delete rows cleanly through high-level API; leave extras
    # only when the target table already has more rows than requested.
    return table


def main():
    path = find_report()
    doc = Document(path)

    # Cover and front-matter callouts.
    set_cell(
        doc.tables[0].cell(0, 0),
        "核心成果　本研究以 MLPClassifier 作為期末主模型，經 GridSearchCV 找到最佳神經網路設定，"
        "並以 OOF F2-score 選出分類門檻約 0.245。測試集 Recall 達 0.78，"
        "False Negative 由 33 人降至 11 人，PR-AUC 為 0.2736，高於 Logistic Regression 基準模型的 0.2517。",
        size=10.5,
        bold=True,
    )
    set_cell(
        doc.tables[1].cell(0, 0),
        "關鍵字　中風預測、神經網路、MLPClassifier、Multilayer Perceptron、類別不平衡、"
        "GridSearchCV、PR-AUC、OOF F2-score、分類門檻",
        size=10.5,
        bold=True,
    )
    set_cell(
        doc.tables[2].cell(0, 0),
        "與期中報告的銜接　期中流程使用 BMI 平均值填補、pandas One-Hot Encoding、MinMaxScaler，"
        "並比較 Logistic Regression、Random Forest 與 KNN。期末延伸改用可避免交叉驗證洩漏的 Pipeline，"
        "以中位數補值、OneHotEncoder 與 StandardScaler 處理資料，並依期末要求以 MLPClassifier 神經網路作為主模型；"
        "Logistic Regression 僅保留為 benchmark / baseline / 可解釋基準模型。",
        size=10.5,
        bold=True,
    )
    set_cell(
        doc.tables[3].cell(0, 0),
        "方法意義　GridSearchCV 決定 MLP 神經網路如何學習；OOF F2 threshold tuning 則決定模型機率如何轉換為分類結果。"
        "兩者都只使用訓練資料完成，測試集保留到最後評估，以避免資料洩漏。",
        size=10.5,
        bold=True,
    )

    # Paragraph replacements by stable content snippets.
    replacements = {
        "統一尺度可避免 age、avg_glucose_level 與 bmi 的數值範圍差異影響 Logistic Regression。":
            "統一尺度可避免 age、avg_glucose_level 與 bmi 的數值範圍差異影響模型訓練；這對 MLP 神經網路特別重要，也讓 LR benchmark 的比較維持在相同前處理基礎上。",
        "圖 4　GridSearchCV 各參數組合的五折交叉驗證 PR-AUC":
            "圖 4　MLP GridSearchCV 與模型比較結果",
        "圖 5　不同分類門檻下的 Precision、Recall 與 F2-score":
            "圖 5　MLP OOF F2 threshold tuning 與門檻選擇概念",
        "圖 6　預設門檻與調整後門檻的測試集混淆矩陣":
            "圖 6　MLP 預設門檻與調整後門檻的測試集混淆矩陣",
        "圖 7　分類門檻調整前後的核心指標與誤判取捨":
            "圖 7　MLP threshold tuning 後的核心指標與誤判取捨",
        "分類問題依老師規定以 Precision、F1-score 與混淆矩陣為主要結果，並額外呈現 Recall、F2-score、Accuracy 與 PR-AUC。調整門檻後 F1-score由 0.3051 降至 0.2676，原因是 Precision 的下降幅度大於 Recall 的提升對 F1 的貢獻；但更重視 Recall 的 F2-score 則由 0.3358 提升至 0.4454。PR-AUC 在兩個門檻下皆為 0.2517，因為 PR-AUC 是由完整預測機率排序計算，不會因單一分類門檻而改變。Accuracy、Precision、Recall 與 F2-score 則依賴最終類別，因此會隨門檻調整。這也說明評估模型時，必須區分「機率排序能力」與「特定門檻下的分類結果」。":
            "分類問題依老師規定以 Precision、F1-score 與混淆矩陣為主要結果，並額外呈現 Recall、F2-score、Accuracy 與 PR-AUC。以 MLP 為主模型時，調整門檻後 Recall 由 0.3400 提升至 0.7800，False Negative 由 33 降至 11，F2-score 由 0.3148 提升至 0.4295；Precision 則由 0.2429 降至 0.1535，False Positive 由 53 增加至 215。PR-AUC 在兩個門檻下皆為 0.2736，因為 PR-AUC 是由完整預測機率排序計算，不會因單一分類門檻而改變。這也說明評估模型時，必須區分「機率排序能力」與「特定門檻下的分類結果」。",
        "門檻調整後 Precision 降至 0.1606，表示模型判定為中風風險者中，真正中風者約占 16%。False Positive 增加至 209 人，可能造成額外檢查、醫療資源負擔與個案焦慮。因此，0.268 並非唯一正確門檻，而是基於 F2-score 與「降低漏判優先」假設得到的選擇。實際部署時應由醫療成本與風險政策共同決定門檻。":
            "",
    }

    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements:
            set_paragraph(p, replacements[t], size=11)
        elif "統一尺度可避免 age、avg_glucose_level 與 bmi 的數值範圍差異影響 Logistic Regression。" in t:
            set_paragraph(
                p,
                t.replace(
                    "統一尺度可避免 age、avg_glucose_level 與 bmi 的數值範圍差異影響 Logistic Regression。",
                    "統一尺度可避免 age、avg_glucose_level 與 bmi 的數值範圍差異影響模型訓練；這對 MLP 神經網路特別重要，也讓 LR benchmark 的比較維持在相同前處理基礎上。",
                ),
                size=11,
            )
        elif "略低於 Logistic Regression 調整門檻後的 Recall 0.80、F2-score 0.4454 與 False Negative 10。" in t:
            set_paragraph(
                p,
                t.replace(
                    "略低於 Logistic Regression 調整門檻後的 Recall 0.80、F2-score 0.4454 與 False Negative 10。",
                    "略低於 Logistic Regression benchmark 調整門檻後的 Recall 0.80、F2-score 0.4454、False Negative 10 與 False Positive 209。",
                ),
                size=11,
            )

    # Main MLP result table in section 8.2.
    table = doc.tables[4]
    normalize_table(table, 11, 4)
    mlp_rows = [
        ["評估指標", "預設門檻 0.500", "調整門檻 0.245", "變化／說明"],
        ["Accuracy", "0.9159", "0.7789", "-0.1370"],
        ["Precision", "0.2429", "0.1535", "-0.0894"],
        ["Recall", "0.3400", "0.7800", "+0.4400"],
        ["F1-score", "0.2833", "0.2566", "-0.0268"],
        ["F2-score", "0.3148", "0.4295", "+0.1147"],
        ["PR-AUC", "0.2736", "0.2736", "不受 threshold 影響"],
        ["True Positive", "17", "39", "+22"],
        ["False Negative", "33", "11", "-22"],
        ["False Positive", "53", "215", "+162"],
        ["定位", "漏判較多", "召回導向", "MLP 主模型結果"],
    ]
    for r, row in enumerate(mlp_rows):
        for c, value in enumerate(row):
            set_cell(table.cell(r, c), value, size=10.2, bold=(r == 0))

    # Keep the comparison table, but make LR role explicit.
    table = doc.tables[6]
    comparison_rows = [
        ["模型定位", "Recall", "F2-score", "FN / FP / PR-AUC"],
        ["MLP 主模型 + threshold 0.245", "0.78", "0.4295", "FN 11 / FP 215 / PR-AUC 0.2736"],
        ["LR benchmark + threshold 0.268", "0.80", "0.4454", "FN 10 / FP 209 / PR-AUC 0.2517"],
        ["判讀", "LR 略高", "LR 略高", "MLP 的 PR-AUC 較高"],
    ]
    for r, row in enumerate(comparison_rows):
        for c, value in enumerate(row):
            set_cell(table.cell(r, c), value, size=10.2, bold=(r == 0))

    # Final conclusion callout should not sound LR-centric.
    set_cell(
        doc.tables[7].cell(0, 0),
        "最終結論　本研究以 MLP 神經網路作為期末主模型。MLP 在 PR-AUC 上高於 LR benchmark，代表風險排序能力有提升；"
        "但 Recall 與 F2 仍略低於 LR，因此本研究不宣稱神經網路全面優於 Logistic Regression，而是誠實呈現它在不平衡中風資料上的優勢與限制。",
        size=10.5,
        bold=True,
    )

    out = path
    try:
        doc.save(out)
        print(f"Updated {out}")
    except PermissionError:
        out = path.with_name(path.stem + "_修正版.docx")
        doc.save(out)
        print(f"Original locked; saved {out}")


if __name__ == "__main__":
    main()
